#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создаёт оформленный Word-шаблон с разделами для всех блоков сайта.
Запуск: python scripts/create-word-template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "WORKOUT_Методические_материалы.docx"

RED = RGBColor(255, 45, 45)
BLACK = RGBColor(20, 20, 24)
GRAY = RGBColor(110, 110, 120)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RED if level == 1 else BLACK
        run.font.name = "Arial"
        run.font.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = BLACK
    return p


def add_meta(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.italic = True
    return p


def main():
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    logo = ROOT / "logo.jpg"
    if logo.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(logo), width=Inches(1.4))
        except Exception:
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run("WORKOUT")
            r.bold = True
            r.font.size = Pt(28)
            r.font.color.rgb = RED

    add_meta(doc, "WORKOUT SPORT CENTER · ЗДОРОВОЕ ПОКОЛЕНИЕ · 2016")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Методические материалы")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RED
    r.font.name = "Arial"

    doc.add_paragraph()

    sections = [
        (
            "1. Организация тренировочного процесса",
            [
                "В наших центрах тренеры сталкиваются с такой трудностью, как разные физические возможности детей в одной группе, ведь при наборе в группу не учитывается физическая подготовка. Поэтому мы создали «Тренировочный год», где расписали под каждый месяц определённую дисциплину занятий.",
                "График тренировочного года расписан так, что у всех занимающихся в год будет 4 раза каждая дисциплина (с учётом физических возможностей).",
                "ПРИМЕР: Январь — контрольная тренировка + динамика. Февраль — статика. Март — ОФП (подготовка к контрольной). И так каждые 3 месяца.",
                "Для удобства разработаны карточки с тренировочными программами для разного уровня ОФП и статики.",
                "Для новых клиентов — правила пробной тренировки и расписание на 8 тренировок с разными темами.",
            ],
        ),
        (
            "2. Подготовительный уровень",
            [
                "Упражнения данного уровня подойдут для маленьких детей, для детей с лишним весом и для тех, кто не развит физически.",
                "Основная задача тренера на подготовительном этапе — научить воспитанников выполнять все назначаемые упражнения с правильной техникой.",
                "Видео: положите файлы в папку video/ (см. README — имена ПОДТЯГИВАНИЯ ПОДГОТОВИТЕЛЬНЫЙ УРОВЕРЬ 1.mp4 и т.д.)",
            ],
        ),
        (
            "3. Разминка и разогрев",
            [
                "Разминку и растяжку проводит тренер или ученик, который знает порядок упражнений.",
                "10–15 минут: общебеговые, суставная гимнастика, динамическая растяжка.",
            ],
        ),
        (
            "4. Продвинутый уровень",
            [
                "Упражнения продвинутого уровня — для атлетов, уверенно выполняющих средний уровень.",
                "Сложные элементы: выходы силой, планш, флажок, комбинации.",
            ],
        ),
        (
            "5. Рекомендации по тренировочному процессу",
            [
                "Упражнения разных уровней — примеры и ориентиры методики.",
                "Адаптируйте нагрузку под группу. Проводите контрольные тренировки.",
            ],
        ),
        (
            "6. Средний уровень",
            [
                "Для атлетов, уверенно выполняющих начальный уровень.",
                "Увеличение объёма, новые вариации подтягиваний и отжиманий.",
            ],
        ),
        (
            "7. Начальный уровень",
            [
                "Для атлетов, умеющих подтягиваться и отжиматься в базовом объёме.",
                "Закрепление техники, сила хвата, стабильность корпуса.",
            ],
        ),
        (
            "8. Разминка и разогрев (дополнительно)",
            [
                "Специальные упражнения на кисти, плечи и локти перед турником.",
                "Завершение: 2–3 подхода лёгких подтягиваний или висов.",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        add_heading(doc, heading, level=1)
        for para in paragraphs:
            add_body(doc, para)
        doc.add_paragraph()

    add_meta(
        doc,
        "После редактирования запустите: python scripts/build-content.py",
    )

    doc.save(OUT)
    print(f"Создан файл: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
